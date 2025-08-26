# app_streamlit_mcp_tables.py
import streamlit as st
import os
import asyncio
from pistachio_mcp_agent import pistachio_query, append_uploaded_files_to_index, INDEX, DATA_DIR, PERSIST_DIR
from io import BytesIO
from docx import Document
import pdfplumber
import html
from llama_index.embeddings.ollama import OllamaEmbedding
from llama_index.core import Settings



import asyncio

def run_async(coro):
    try:
        loop = asyncio.get_running_loop()  # check if already running
    except RuntimeError:
        return asyncio.run(coro)  # no loop, safe
    else:
        return loop.run_until_complete(coro)  # reuse running loop


st.set_page_config(page_title="Pistachio Assistant", page_icon="🌰", layout="wide")
st.title("🌰 Pistachio Growing Assistant (Local RAG + MCP)")
st.caption("Runs locally with Ollama + LlamaIndex. Upload documents and ask questions.")

# ------------------- Upload Documents -------------------
st.sidebar.header("📁 Upload Documents")
uploaded_files = st.sidebar.file_uploader(
    "Add PDFs, Word docs, or text files",
    type=["pdf", "txt", "md", "docx"],
    accept_multiple_files=True
)
add_btn = st.sidebar.button("Add to Index")

# if uploaded_files and add_btn:
#     with st.spinner("Indexing uploaded files..."):
#         added_count = append_uploaded_files_to_index(INDEX, uploaded_files)
#     st.sidebar.success(f"Indexed {added_count} new document(s).")

# ------------------- Query UI -------------------
st.subheader("Ask a question about pistachios")

with st.form(key="query_form"):
    user_query = st.text_input("e.g., What soil pH is best for pistachio trees?", key="user_query")
    submit_btn = st.form_submit_button("Ask")

if submit_btn and user_query.strip():
    with st.spinner("Thinking locally..."):
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        #loop = asyncio.get_event_loop()
        #answer = run_until_complete(pistachio_query(user_query))
        answer = run_async(pistachio_query(user_query))


    st.markdown("### Answer")
    st.write(answer)

    # ------------------- Sources & Highlights -------------------

    # # Local embedding model
    # embed_model = OllamaEmbedding(model_name="nomic-embed-text")
    # set_global_embedding(embed_model)

    # # Local LLM (e.g. llama2 or mistral)
    # llm = Ollama(model="llama3.2")
    # set_global_llm(llm)

 
    # embed_model_name = os.getenv("EMBED_MODEL", "nomic-embed-text")  # 'local' not directly supported
    # base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")

    # Settings.embed_model = OllamaEmbedding(model_name=embed_model_name, base_url=base_url)


    qe = INDEX.as_query_engine()
    response = qe.query(user_query)

    if hasattr(response, "source_nodes") and response.source_nodes:
        st.markdown("#### Sources & Relevant Passages")
        query_terms = [term.lower() for term in user_query.split() if term.strip()]

        for i, sn in enumerate(response.source_nodes, start=1):
            meta = sn.node.metadata or {}
            fname = meta.get("file_name") or meta.get("filename") or meta.get("source") or "document"
            score = sn.score if sn.score is not None else 0
            color = "#a6cee3" if score > 0.8 else "#b2df8a" if score > 0.5 else "#fb9a99"

            content_preview = sn.node.get_content()
            lines = content_preview.split("\n")

            # Highlight query terms in snippet
            highlighted_lines = []
            for line in lines[:20]:
                hl_line = line
                for term in query_terms:
                    if term in line.lower():
                        hl_line = hl_line.replace(term, f"<span style='color:red;font-weight:bold'>{term}</span>")
                highlighted_lines.append(html.escape(hl_line))
            highlighted_text = "<br>".join(highlighted_lines)

            with st.expander(f"{i}. {fname} (score {score:.3f})", expanded=False):
                # Snippet preview
                st.markdown(f"<div style='background-color:{color};padding:5px;border-radius:5px'>{highlighted_text}</div>", unsafe_allow_html=True)

                # ------------------- Rich Document Preview with Tables -------------------
                doc_path = meta.get("file_path") or meta.get("source") or None
                if doc_path and os.path.exists(doc_path):
                    ext = os.path.splitext(doc_path)[-1].lower()
                    html_content = ""

                    if ext == ".pdf":
                        # PDF: basic text + try extracting tables
                        with pdfplumber.open(doc_path) as pdf:
                            pages_html = []
                            for page in pdf.pages:
                                text = html.escape(page.extract_text() or "").replace("\n", "<br>")
                                tables_html = ""
                                for table in page.extract_tables():
                                    tables_html += "<table border='1' style='border-collapse:collapse;margin:5px;'>"
                                    for row in table:
                                        tables_html += "<tr>" + "".join(f"<td style='padding:3px'>{html.escape(str(cell or ''))}</td>" for cell in row) + "</tr>"
                                    tables_html += "</table>"
                                pages_html.append(text + tables_html)
                            html_content = "<hr>".join(pages_html)

                    elif ext == ".docx":
                        doc = Document(doc_path)
                        html_parts = []
                        for p in doc.paragraphs:
                            text = html.escape(p.text)
                            # Bold/Italic
                            if p.runs:
                                for r in p.runs:
                                    if r.bold:
                                        text = text.replace(html.escape(r.text), f"<b>{html.escape(r.text)}</b>")
                                    if r.italic:
                                        text = text.replace(html.escape(r.text), f"<i>{html.escape(r.text)}</i>")
                            # Bullet / numbered lists
                            if p.style.name.startswith("List"):
                                html_parts.append(f"<li>{text}</li>")
                            else:
                                html_parts.append(f"<p>{text}</p>")
                        # Tables
                        for table in doc.tables:
                            table_html = "<table border='1' style='border-collapse:collapse;margin:5px;'>"
                            for row in table.rows:
                                table_html += "<tr>" + "".join(f"<td style='padding:3px'>{html.escape(cell.text)}</td>" for cell in row.cells) + "</tr>"
                            table_html += "</table>"
                            html_parts.append(table_html)
                        html_content = "".join(html_parts)

                    else:
                        with open(doc_path, "r", encoding="utf-8", errors="ignore") as f:
                            html_content = "<p>" + html.escape(f.read()).replace("\n", "<br>") + "</p>"

                    st.markdown(
                        f"<div style='height:500px;overflow-y:scroll;padding:10px;border:1px solid #ccc;border-radius:5px;background-color:#f8f8f8'>{html_content}</div>",
                        unsafe_allow_html=True
                    )
                else:
                    st.write("Full document not available locally.")
